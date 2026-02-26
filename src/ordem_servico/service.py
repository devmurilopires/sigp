import os
import unicodedata
from datetime import datetime
from docx import Document
from docx.shared import Inches
from src.ordem_servico.repository import OSRepository

# Tenta puxar o utils da raiz ou da pasta shared
try:
    from src.shared.utils import resource_path
except ImportError:
    try:
        from src.shared.utils import resource_path
    except ImportError:
        def resource_path(path): return path

class OSService:
    def __init__(self):
        # Conecta o serviço ao repositório (banco de dados)
        self.repo = OSRepository()

    # =========================================================
    # REGRAS DE NEGÓCIO E VALIDAÇÕES
    # =========================================================
    def normalizar(self, texto: str) -> str:
        """Remove acentos e deixa a string em caixa alta para padronização no banco."""
        if not texto: return ""
        nfkd = unicodedata.normalize('NFKD', texto)
        sem_acentos = "".join(c for c in nfkd if not unicodedata.combining(c))
        return sem_acentos.upper()

    def consultar_endereco(self, id_procurado):
        """Busca os dados do endereço no repositório para preencher a tela."""
        return self.repo.buscar_endereco_por_id(id_procurado)

    def obter_historico_formatado(self, id_procurado):
        """Busca o histórico de OS desse ID e formata em texto para mostrar no popup."""
        historico = self.repo.buscar_historico_os(id_procurado)
        if not historico:
            return "Nenhuma movimentação de Ordem de Serviço encontrada para este ID."
        
        texto = ""
        for r in historico:
            texto += f"OS: {r[0]} | Data: {r[1]} | Tipo: {r[2]} | Item: {r[3]}\nEndereço: {r[4]} - Bairro: {r[5]}\nCriado por: {r[6]}\n"
            texto += "-" * 50 + "\n"
        return texto

    # =========================================================
    # ORQUESTRAÇÃO PRINCIPAL (GERAÇÃO SEGURA - DB PRIMEIRO)
    # =========================================================
    def processar_criacao_os(self, descricoes_acumuladas, pasta_escolhida, modelo_escolhido, tipo_os, tipo_item, form_dados, usuario_logado, origem_demanda):
        if not descricoes_acumuladas:
            return False, "Adicione pelo menos um item (descrição) na lista antes de gerar a OS."

        # ---> NOVO CAMINHO DINÂMICO E INTELIGENTE DA REDE <---
        ano_atual = datetime.now().strftime('%Y')
        raiz_rede = r"C:\Users\sousa\OneDrive\Documentos\ARQUIVOS SIGP - SIGA - SPR"
        
        # Só bloqueia se o Servidor/Rede estiver fora do ar. As pastas do ano ele cria sozinho!
        if not os.path.exists(raiz_rede):
            return False, f"A raiz da rede não está acessível no momento. Verifique a conexão:\n{raiz_rede}"

        if pasta_escolhida == "URBMIDIA":
            pasta_base = rf"{raiz_rede}\SIGP\{ano_atual}\ORDENS DE SERVICO\URBMIDIA"
        else:
            pasta_base = rf"{raiz_rede}\SIGP\{ano_atual}\ORDENS DE SERVICO\PROXIMA PARADA"

        ids_unicos = list(set([d["id"] for d in descricoes_acumuladas]))
        ids_formatado = "-".join(ids_unicos)
        id_principal = descricoes_acumuladas[0]["id"]

        for id_atual in ids_unicos:
            dados_id = self.repo.buscar_endereco_por_id(id_atual)
            try:
                if not dados_id:
                    self.repo.cadastrar_endereco(
                        id_atual, form_dados['endereco'], form_dados['numero'], 
                        form_dados['bairro'], form_dados['complemento'], usuario_logado
                    )
                else:
                    reativar = dados_id["status"].strip().upper() == "INATIVO"
                    self.repo.atualizar_endereco(
                        id_atual, form_dados['endereco'], form_dados['numero'], 
                        form_dados['bairro'], form_dados['complemento'], usuario_logado, reativar=reativar
                    )
            except Exception as e:
                return False, f"Erro ao gerenciar endereços no banco:\n{str(e)}"

        numero_os = self.repo.obter_proximo_numero_os(pasta_escolhida, ano_atual)
        data_str = datetime.now().strftime("%d/%m/%Y")

        endereco_completo = descricoes_acumuladas[0]["descricao"].split(" NA ")[-1].split(",")[0].strip()
        bairro_str = form_dados['bairro']
        complemento_str = form_dados['complemento']
        
        try:
            if "BAIRRO" in descricoes_acumuladas[0]["descricao"]:
                bairro_str = descricoes_acumuladas[0]["descricao"].split("BAIRRO")[-1].split(",")[0].strip()
            if "-" in bairro_str:
                partes = bairro_str.split("-", 1)
                bairro_str = partes[0].strip()
                complemento_str = partes[1].strip()
        except:
            pass 

        tipo_os_up = str(tipo_os).strip().upper() if tipo_os else ""
        tipo_item_up = str(tipo_item).strip().upper() if tipo_item else ""

        dados_salvar_os = (
            numero_os, data_str, id_principal, ids_formatado,
            tipo_os_up, self.normalizar(tipo_os_up),
            tipo_item_up, self.normalizar(tipo_item_up),
            endereco_completo, bairro_str, self.normalizar(bairro_str),
            complemento_str, "\n".join([item["descricao"] for item in descricoes_acumuladas]),
            usuario_logado, pasta_escolhida, origem_demanda 
        )

        try:
            self.repo.salvar_os(dados_salvar_os)
        except Exception as e:
            return False, f"Erro Crítico! A OS NÃO foi gerada pois houve falha no Banco de Dados:\n{str(e)}"

        nome_pasta = f"{numero_os:03d}-{datetime.now().strftime('%m')}-{ano_atual}-ID{'-'.join(ids_unicos) if ids_unicos else 'EMERGENCIA'}"
        caminho_pasta = os.path.join(pasta_base, nome_pasta)
        nome_arquivo = f"O.S {numero_os:03d}-{ano_atual}-ID{'-'.join(ids_unicos) if ids_unicos else 'EMERGENCIA'}.docx"
        destino_docx = os.path.join(caminho_pasta, nome_arquivo)

        try:
            # ---> O MÁGICO os.makedirs AQUI VAI CRIAR A PASTA DO ANO CASO NÃO EXISTA
            os.makedirs(caminho_pasta, exist_ok=True)
            caminho_modelo = resource_path(modelo_escolhido)
            self._gerar_documento_modelo(caminho_modelo, destino_docx, numero_os, data_str, id_principal, descricoes_acumuladas)
            
            return True, f"Ordem de Serviço Nº {numero_os:03d} criada e registrada com sucesso!\nSalva em:\n{destino_docx}"
            
        except Exception as e:
            return False, f"Atenção: A OS foi registrada no banco, mas houve falha ao gerar o documento Word na rede:\n{e}"

    # =========================================================
    # MANIPULAÇÃO DO WORD (DOCX)
    # =========================================================
    def _gerar_documento_modelo(self, modelo_path, destino_path, numero_os, data_str, id_texto, descricoes):
        """Abre o modelo do Word, substitui as tags e gera a tabela de descrições."""
        doc = Document(modelo_path)
        mapeamento = {
            "{{NUMERO_OS}}": f"{numero_os:03d}",
            "{{DATA}}": data_str,
            "{{ID}}": id_texto if id_texto.strip() else "-"
        }
        
        for paragrafo in doc.paragraphs:
            texto_original = "".join(run.text for run in paragrafo.runs)
            novo_texto = texto_original
            for chave, valor in mapeamento.items():
                if chave in novo_texto:
                    novo_texto = novo_texto.replace(chave, valor)
            if novo_texto != texto_original:
                for run in paragrafo.runs: run.text = ""
                if paragrafo.runs: paragrafo.runs[0].text = novo_texto

        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for paragrafo in celula.paragraphs:
                        texto_original = "".join(run.text for run in paragrafo.runs)
                        novo_texto = texto_original
                        for chave, valor in mapeamento.items():
                            if chave in novo_texto:
                                novo_texto = novo_texto.replace(chave, valor)
                        if novo_texto != texto_original:
                            for run in paragrafo.runs: run.text = ""
                            if paragrafo.runs: paragrafo.runs[0].text = novo_texto

        for paragrafo in doc.paragraphs:
            if "{{DESCRICAO}}" in paragrafo.text:
                p = paragrafo._element
                parent = p.getparent()
                
                tabela = doc.add_table(rows=1, cols=2)
                tabela.style = 'Table Grid'
                tabela.columns[0].width = Inches(1.0)
                tabela.columns[1].width = Inches(5.0)
                
                hdr_cells = tabela.rows[0].cells
                hdr_cells[0].text = 'ID'
                hdr_cells[1].text = 'DESCRIÇÃO'
                
                for item in descricoes:
                    row_cells = tabela.add_row().cells
                    row_cells[0].text = item['id']
                    row_cells[1].text = item['descricao']
                
                p.addnext(tabela._element)
                parent.remove(p)
                break
        
        doc.save(destino_path)