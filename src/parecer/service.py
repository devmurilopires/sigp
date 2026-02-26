import os
from datetime import datetime
from docx import Document
from src.parecer.repository import ParecerRepository

try:
    from src.shared.utils import resource_path
except ImportError:
    def resource_path(path): return path

class ParecerService:
    def __init__(self):
        self.repo = ParecerRepository()

    def processar_geracao_parecer(self, dados_form, ids_list, usuario_logado):
        if not ids_list:
            return False, "Adicione ao menos um ID antes de gerar o parecer."

        # Extrai os dados do formulário
        origem = dados_form['origem']
        tipo_parecer = dados_form['tipo']
        processo = dados_form['processo'] or "-"
        assunto = dados_form['assunto'] or "-"
        solicitante = dados_form['solicitante'] or "-"
        tipo_exec = dados_form['tipo_execucao'] or "-"
        item = dados_form['item'] or "-"
        endereco = dados_form['endereco'] or "-"
        motivo = dados_form['motivo'] if tipo_parecer == "Indeferido" else "-"
        quantidade = dados_form['quantidade'] or "-"
        
        ids_joined = ", ".join(ids_list)
        data_atual = datetime.now()
        ano = data_atual.year
        data_str = data_atual.strftime("%d/%m/%Y")

        # 1. Ajusta o Plural do Item
        quantidade_normalizada = quantidade.lower().strip()
        plurais = {
            "Abrigo Metálico": "Abrigos Metálicos",
            "Placa/Barrote": "Placas/Barrote",
            "Placa/Poste": "Placas/Poste",
            "Parada Segura": "Paradas Seguras",
            "Abrigo Concreto": "Abrigos Concretos"
        }
        if not (quantidade_normalizada.startswith("um") or quantidade_normalizada.startswith("uma")):
            item = plurais.get(item, item)

        # 2. Pega o número do parecer
        try:
            numero = self.repo.obter_proximo_numero(ano)
        except Exception as e:
            return False, str(e)

        # 3. Prepara Caminhos do Arquivo Word
        modelo = resource_path(os.path.join("dados", "modelo_deferido.docx")) if tipo_parecer == "Deferido" else resource_path(os.path.join("dados", "modelo_indeferido.docx"))
        
        if not os.path.exists(modelo):
            return False, f"Modelo Word não encontrado em: {modelo}"

        # ---> NOVO CAMINHO DINÂMICO E INTELIGENTE DA REDE <---
        raiz_rede = r"\\172.20.0.57\dados\DIPLA\ARQUIVOS SIGP - SIGA - SPR"
        
        # Só bloqueia se o Servidor/Rede estiver fora do ar.
        if not os.path.exists(raiz_rede):
            return False, f"A raiz da rede não está acessível no momento. Verifique a conexão:\n{raiz_rede}"

        pasta_base = rf"{raiz_rede}\SIGP\{ano}\PARECERES TECNICOS"
        pasta_saida = os.path.join(pasta_base, tipo_parecer.upper())
        
        nome_arquivo = f"Parecer_{numero:03d}_{ano}_{tipo_parecer}.docx"
        caminho_arquivo = os.path.join(pasta_saida, nome_arquivo)

        # Prepara os dados para o banco (Com a ORIGEM no final)
        dados_banco = (
            numero, ano, data_atual.date(), tipo_parecer.upper(), processo, 
            assunto, ids_joined, tipo_exec, item, endereco, 
            solicitante, motivo if tipo_parecer == "Indeferido" else None, 
            quantidade, caminho_arquivo, usuario_logado, origem 
        )

        # =========================================================================
        # 4. GERAÇÃO SEGURA (BANCO DE DADOS PRIMEIRO)
        # =========================================================================
        try:
            self.repo.salvar_parecer(dados_banco)
        except Exception as e:
            return False, f"Erro Crítico! O Parecer NÃO foi gerado pois houve falha no Banco de Dados:\n{str(e)}"

        # =========================================================================
        # 5. SE O BANCO DEU CERTO -> GERA A PASTA E O WORD
        # =========================================================================
        try:
            # ---> O MÁGICO os.makedirs AQUI VAI CRIAR A PASTA DO ANO CASO NÃO EXISTA
            os.makedirs(pasta_saida, exist_ok=True)
            self._gerar_documento_word(modelo, caminho_arquivo, {
                "{{NUM_PARECER}}": f"{numero:03d}",
                "{{DATA}}": data_str,
                "{{PROCESSO}}": processo,
                "{{ASSUNTO}}": assunto,
                "{{SOLICITANTE}}": solicitante,
                "{{ID}}": ids_joined,
                "{{TIPO}}": tipo_exec,
                "{{ITEM}}": item,
                "{{ENDERECO}}": endereco,
                "{{MOTIVO}}": motivo,
                "{{QUANTIDADE}}": quantidade
            })
            
            return True, f"Parecer {numero:03d}/{ano} criado e registrado com sucesso!\nSalvo em:\n{caminho_arquivo}"
            
        except Exception as e:
            return False, f"Atenção: O Parecer foi registrado no banco, mas houve falha ao gerar o documento Word:\n{e}"

    # =========================================================
    # MANIPULAÇÃO DO WORD (DOCX)
    # =========================================================
    def _gerar_documento_word(self, modelo_path, destino_path, tags):
        doc = Document(modelo_path)
        
        for p in doc.paragraphs:
            texto_completo = "".join(run.text for run in p.runs)
            modificado = False
            for tag, val in tags.items():
                if tag in texto_completo:
                    texto_completo = texto_completo.replace(tag, val)
                    modificado = True
            
            if modificado:
                for run in p.runs: run.text = ""
                if p.runs: p.runs[0].text = texto_completo

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        texto_completo = "".join(run.text for run in p.runs)
                        modificado = False
                        for tag, val in tags.items():
                            if tag in texto_completo:
                                texto_completo = texto_completo.replace(tag, val)
                                modificado = True
                        if modificado:
                            for run in p.runs: run.text = ""
                            if p.runs: p.runs[0].text = texto_completo
        
        doc.save(destino_path)